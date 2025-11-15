import os
from flask import Flask, render_template, request, url_for, redirect, send_file, flash
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from PIL import Image

# --- Configuration ---
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
ALLOWED_IMG = {"png", "jpg", "jpeg", "gif"}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app = Flask(__name__)
app.secret_key = "replace-with-a-secure-key"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 64 * 1024 * 1024

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_IMG

def save_file(fileobj):
    if not fileobj:
        return None
    filename = secure_filename(fileobj.filename)
    path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    fileobj.save(path)
    return path

# ---- Helpers to build docx ----
def build_docx(path, data, invitation_path, photo_paths):
    doc = Document()

    # Setup header with logo and college name (applies to all pages)
    section = doc.sections[0]
    header = section.header
    hdr_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hdr_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = hdr_p.add_run()
    try:
        run.add_picture(os.path.join("static", "sec_logo.png"), width=Inches(1.6))
    except Exception:
        pass
    hdr_p.add_run("\nSURYA ENGINEERING COLLEGE, ERODE").bold = True

    # Page 1: Circular (matching uploaded sample layout)
    # Reference and Date row (left blank and right aligned like sample)
    doc.add_paragraph()  # spacing
    p_ref = doc.add_paragraph()
    p_ref.add_run(f"Ref: {data['ref_no']}\t\t\tDate: {data['date_ref']}\n").bold = True

    doc.add_paragraph("\nCIRCULAR\n", style='Intense Quote')

    # Invitation/circular image or text
    if invitation_path:
        try:
            doc.add_picture(invitation_path, width=Inches(6))
        except Exception:
            doc.add_paragraph(data.get("circular_text", ""))
    else:
        doc.add_paragraph(data.get("circular_text", ""))

    # Copy to / signatures area like sample
    doc.add_paragraph("\n\nHoD\t\t\t\tPRINCIPAL")
    doc.add_paragraph("\nCopy to\nAll HoD's\nNotice board\nTo be read in all class rooms")

    doc.add_page_break()

    # Page 2: Report header (department + report title)
    doc.add_paragraph(f"DEPARTMENT OF {data.get('department','')}", style='Title')
    doc.add_paragraph("REPORT OF THE EVENT", style='Heading 1')

    # Report details table (mimic sample table)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List Accent 1'
    rows = [
        ("Name of the activity", data.get("title", "")),
        ("Department", data.get("department", "")),
        ("Academic year", data.get("academic_year", "")),
        ("Date of activity conducted", data.get("date_activity", "")),
        ("Organized by", data.get("organized_by", "")),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    # Add structured sections (Overview, Session Highlights, Outcomes, Conclusion)
    doc.add_paragraph("\n")
    doc.add_paragraph("Workshop Overview", style='Heading 2')
    doc.add_paragraph(data.get("overview", ""))

    if data.get("day1"):
        doc.add_paragraph("\nDay 1: Fundamentals and Demonstrations", style='Heading 3')
        doc.add_paragraph(data.get("day1", ""))
    if data.get("day2"):
        doc.add_paragraph("\nDay 2: Advanced Troubleshooting and Maintenance", style='Heading 3')
        doc.add_paragraph(data.get("day2", ""))

    if data.get("learning_outcomes"):
        doc.add_paragraph("\nLearning Outcomes", style='Heading 2')
        doc.add_paragraph(data.get("learning_outcomes"))

    doc.add_paragraph("\nConclusion", style='Heading 2')
    doc.add_paragraph(data.get("conclusion", ""))

    doc.add_paragraph("\n\nFunction was inaugurated by our resource person\n\n")
    # Leave space then signatures as in sample
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Light List Accent 1'
    table2.rows[0].cells[0].text = "\n\nHoD"
    table2.rows[0].cells[1].text = "\n\nPRINCIPAL"

    doc.add_page_break()

    # Page 3: Photos layout (grid)
    doc.add_paragraph("Event Photographs", style='Heading 1')
    if photo_paths:
        # Insert images two per row
        for i, p in enumerate(photo_paths):
            try:
                # Create a 2-cell row table for each pair
                if i % 2 == 0:
                    row_cells = doc.add_table(rows=1, cols=2).rows[0].cells
                run = row_cells[i % 2].paragraphs[0].add_run()
                run.add_picture(p, width=Inches(3))
            except Exception:
                # if image insertion fails, continue
                doc.add_paragraph(f"(Unable to add image: {os.path.basename(p)})")
    else:
        doc.add_paragraph("No photos uploaded.")

    doc.save(path)
    return path

# ---- Helpers to build PDF ----
def build_pdf(path, data, invitation_path, photo_paths):
    # Create PDF with header on each page
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    logo_path = os.path.join("static", "sec_logo.png")

    def draw_header():
        # draw logo centered
        try:
            im = Image.open(logo_path)
            iw, ih = im.size
            target_w = 120  # px approximate
            ratio = target_w / iw
            target_h = ih * ratio
            c.drawInlineImage(logo_path, (w/2 - target_w/2), h - 80, width=target_w, height=target_h)
        except Exception:
            pass
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(w/2, h-95, "SURYA ENGINEERING COLLEGE, ERODE")

    # Page 1 - Circular
    draw_header()
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, h-120, f"Ref: {data.get('ref_no','')}")
    c.drawRightString(w-50, h-120, f"Date: {data.get('date_ref','')}")
    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, h-150, "CIRCULAR")

    # invitation image
    y = h-180
    if invitation_path:
        try:
            im = Image.open(invitation_path)
            iw, ih = im.size
            max_w = w - 100
            ratio = max_w / iw
            draw_w = max_w
            draw_h = ih * ratio
            c.drawInlineImage(invitation_path, 50, y-draw_h-10, width=draw_w, height=draw_h)
            y = y - draw_h - 40
        except Exception:
            c.setFont("Helvetica", 10)
            c.drawString(50, y, data.get("circular_text",""))
            y -= 20
    else:
        c.setFont("Helvetica", 10)
        c.drawString(50, y, data.get("circular_text",""))
        y -= 20

    c.drawString(50, 120, "HoD\t\t\t\tPRINCIPAL")
    c.showPage()

    # Page 2 - Report
    draw_header()
    c.setFont("Helvetica-Bold", 13)
    c.drawCentredString(w/2, h-120, f"DEPARTMENT OF {data.get('department','')}")
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(w/2, h-140, "REPORT OF THE EVENT")
    y = h-170
    c.setFont("Helvetica", 10)
    lines = [
        f"Name of the activity: {data.get('title','')}",
        f"Department: {data.get('department','')}",
        f"Academic year: {data.get('academic_year','')}",
        f"Date of activity conducted: {data.get('date_activity','')}",
        f"Organized by: {data.get('organized_by','')}"
    ]
    for L in lines:
        c.drawString(50, y, L)
        y -= 16

    y -= 8
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "Workshop Overview:")
    y -= 16
    c.setFont("Helvetica", 10)
    from textwrap import wrap
    for ln in wrap(data.get("overview",""), 100):
        c.drawString(50, y, ln)
        y -= 14
        if y < 120:
            c.showPage()
            draw_header()
            y = h-120

    # Day1, Day2
    if data.get("day1"):
        y -= 8
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, "Day 1: Fundamentals and Demonstrations")
        y -= 16
        c.setFont("Helvetica", 10)
        for ln in wrap(data.get("day1",""), 100):
            c.drawString(50, y, ln)
            y -= 14
            if y < 120:
                c.showPage()
                draw_header()
                y = h-120

    if data.get("day2"):
        y -= 8
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, "Day 2: Advanced Troubleshooting and Maintenance")
        y -= 16
        c.setFont("Helvetica", 10)
        for ln in wrap(data.get("day2",""), 100):
            c.drawString(50, y, ln)
            y -= 14
            if y < 120:
                c.showPage()
                draw_header()
                y = h-120

    # Learning Outcomes & Conclusion
    if data.get("learning_outcomes"):
        y -= 8
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, "Learning Outcomes:")
        y -= 16
        c.setFont("Helvetica", 10)
        for ln in wrap(data.get("learning_outcomes",""), 100):
            c.drawString(50, y, ln)
            y -= 14
            if y < 120:
                c.showPage()
                draw_header()
                y = h-120

    if data.get("conclusion"):
        y -= 8
        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, "Conclusion:")
        y -= 16
        c.setFont("Helvetica", 10)
        for ln in wrap(data.get("conclusion",""), 100):
            c.drawString(50, y, ln)
            y -= 14
            if y < 120:
                c.showPage()
                draw_header()
                y = h-120

    # Signatures area
    c.setFont("Helvetica", 10)
    c.drawString(50, 100, "HoD")
    c.drawRightString(w-50, 100, "PRINCIPAL")
    c.showPage()

    # Page 3 - Photos grid
    draw_header()
    c.setFont("Helvetica-Bold", 13)
    c.drawCentredString(w/2, h-120, "Event Photographs")
    x_margin = 50
    y_start = h-160
    col_w = (w - 2*x_margin - 20) / 2
    row_h = 120
    x_positions = [x_margin, x_margin + col_w + 20]
    col = 0
    row = 0
    for idx, p in enumerate(photo_paths):
        try:
            im = Image.open(p)
            iw, ih = im.size
            ratio = min(col_w / iw, row_h / ih)
            draw_w = iw * ratio
            draw_h = ih * ratio
            x = x_positions[col] + (col_w - draw_w)/2
            y = y_start - row*(row_h + 20) - draw_h
            c.drawInlineImage(p, x, y, width=draw_w, height=draw_h)
        except Exception:
            pass
        col += 1
        if col > 1:
            col = 0
            row += 1
            if y - row_h < 120:
                c.showPage()
                draw_header()
                y_start = h-160
                row = 0

    c.save()
    return path

# ---- Routes ----
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    # read full form – manual ref number & date_ref included
    data = {
        "ref_no": request.form.get("ref_no", ""),
        "date_ref": request.form.get("date_ref", ""),
        "title": request.form.get("title", ""),
        "department": request.form.get("department", ""),
        "academic_year": request.form.get("academic_year", ""),
        "date_activity": request.form.get("date_activity", ""),
        "organized_by": request.form.get("organized_by", ""),
        "overview": request.form.get("overview", ""),
        "day1": request.form.get("day1", ""),
        "day2": request.form.get("day2", ""),
        "learning_outcomes": request.form.get("learning_outcomes", ""),
        "conclusion": request.form.get("conclusion", ""),
        # extras for circular
        "circular_text": request.form.get("circular_text", "")
    }

    # Save invitation and photos
    invitation_file = request.files.get("invitation")
    invitation_path = None
    if invitation_file and allowed_file(invitation_file.filename):
        invitation_path = save_file(invitation_file)

    photos = request.files.getlist("photos")
    photo_paths = []
    for p in photos:
        if p and allowed_file(p.filename):
            photo_paths.append(save_file(p))

    # Create DOCX preview file in outputs
    base_name = secure_filename((data.get("title") or "event").replace(" ", "_"))
    docx_path = os.path.join(app.config["OUTPUT_FOLDER"], f"{base_name}.docx")
    pdf_path = os.path.join(app.config["OUTPUT_FOLDER"], f"{base_name}.pdf")

    # Build DOCX (used as canonical source)
    try:
        build_docx(docx_path, data, invitation_path, photo_paths)
    except Exception as e:
        flash(f"Error creating DOCX: {e}", "danger")
        return redirect(url_for("index"))

    # Build PDF for download option (but we won't force convert unless user chooses)
    try:
        build_pdf(pdf_path, data, invitation_path, photo_paths)
    except Exception:
        # ignore PDF build errors; we can attempt to build on-demand later
        pass

    # Render a fairly exact preview (type A) using the same data and saved images
    # preview.html will show layout similar to doc format
    return render_template("preview.html",
                           data=data,
                           invitation=invitation_path,
                           photos=photo_paths,
                           docx=os.path.basename(docx_path),
                           pdf=os.path.basename(pdf_path))

@app.route("/download/<fmt>/<filename>")
def download(fmt, filename):
    # fmt either 'docx' or 'pdf'
    fpath = os.path.join(app.config["OUTPUT_FOLDER"], filename)
    if fmt == "pdf":
        # if PDF not exist, try to build from docx data? We'll just return if present.
        if not os.path.exists(fpath):
            return "PDF not available", 404
    if not os.path.exists(fpath):
        return "File not found", 404
    return send_file(fpath, as_attachment=True)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)

